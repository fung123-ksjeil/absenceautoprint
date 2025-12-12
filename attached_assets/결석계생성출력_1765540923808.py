

import os
import sys
# ms docx 패키지 임포트
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# openpyxl 패키지를 임포트합니다.
import openpyxl
import re
import win32com.client as win32
from tkinter import messagebox

###################### 질병결석신고서 생성하는 함수 #####################
def Make_Sick_Absence(doc, input_data):
   
    # 문서의 기본 스타일 설정
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1
    style.font.name = 'Gulim'
    style.font.size = Pt(15)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
    
    # 결석신고서 글자크기 30으로 변경하여 입력
    paragraph_a = doc.add_paragraph()
    paragraph_a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_a = paragraph_a.add_run('\n결  석  신  고  서')
    run_a.font.bold = True
    run_a.font.size = Pt(30)

    # 
    paragraph_b = doc.add_paragraph()
    paragraph_b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_b = paragraph_b.add_run('(질병 결석 처리용)\n')
    run_b.font.bold = True
    run_b.font.size = Pt(15)

    #
    paragraph_c = doc.add_paragraph()
    str_temp = (f'                                                         학번 : {input_data[1]}\n' +
                f'                                                         성명 : {input_data[2]}\n')
    run_c = paragraph_c.add_run(str_temp)

    #
    paragraph_d = doc.add_paragraph()
    paragraph_d.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_spacing_line(paragraph_d, 2)
    str_temp = (f'위 본인은 ( {input_data[5]} )부터 ( {input_data[6]} )까지 ( {input_data[7]} )일간 ( {input_data[4]} )(으)로 인하여 결석하였기에 이를 확인하고 결석신고서를 제출합니다.\n\n')
    run_d = paragraph_d.add_run(str_temp)
    
    #
    paragraph_e = doc.add_paragraph()
    paragraph_e.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT    
    str_temp = (f'{input_data[8]}\n\n담 임 :  {input_data[9]}  (인)\n\n\n\n')
    run_e = paragraph_e.add_run(str_temp)

    #               
    paragraph_f = doc.add_paragraph()
    paragraph_f.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    str_temp = '군  산  제  일  중  학  교  장  귀  하\n'
    run_f = paragraph_f.add_run(str_temp)
    run_f.font.bold = True

    # 문서 경로와 이름
    file_name = f'{current_dir}\\질병결석신고서_{input_data[2]}_{input_data[5]}.docx'
    
    return file_name

#####  줄 간격 변경하는 함수, 특별결석신고서 생성하는 함수에서 호출
def set_paragraph_spacing_line(paragraph, line_spacing):
    paragraph.paragraph_format.line_spacing = line_spacing

#####  문단 뒤에 추가되는 공간을 설정
def set_paragraph_spacing_after(paragraph, space_after):
    paragraph.paragraph_format.space_after = Pt(space_after)

#####  문단 앞에 추가되는 공간을 설정
def set_paragraph_spacing_before(paragraph, space_before):
    paragraph.paragraph_format.space_before = Pt(space_before)


###################### 특별결석신고서 생성하는 함수 ##############ㅍ#####
def Make_Spec_Absence(doc, input_data) :
   
    # 문서의 기본 스타일 설정
    style = doc.styles['Normal']
    style.paragraph_format.line_spacing = 1
    style.font.name = 'Gulim'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Gulim')
    
    # 글자입력 / 가운데 정렬, 글자크기 20
    paragraph_a = doc.add_paragraph()
    paragraph_a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER 
    run_a = paragraph_a.add_run('특 별 결 석 신 고 서')
    run_a.font.underline = True
    run_a.font.bold = True
    run_a.font.size = Pt(20)
    
    # 결제란 이미지 삽입 /  오른쪽 정렬, 윗공백 10설정, 줄간격 3
    paragraph_img  = doc.add_paragraph()
    paragraph_img.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    set_paragraph_spacing_before(paragraph_img, 10)
    set_paragraph_spacing_line(paragraph_img, 3)
    run_img = paragraph_img.add_run()
    
    run_img.add_picture(f'{current_dir}\\삽입이미지.png')
    
    # 글자 입력하기 / 줄 간격 배수 : 1.5, 글자크기 12
    paragraph_b = doc.add_paragraph()
    set_paragraph_spacing_line(paragraph_b, 1.5)
    str_temp =  (f'                                                                             학번 : {input_data[1]}\n' +   
                 f'                                                                             성명 : {input_data[2]}'
                 )
    run_b = paragraph_b.add_run(str_temp)
    run_b.font.size = Pt(12)     

    # 글자 입력하기 / 왼쪽 정렬, 줄 간격 배수 : 1.5, 글자크기 12
    paragraph_c = doc.add_paragraph()
    paragraph_c.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_paragraph_spacing_line(paragraph_c, 2)
    str_temp = (f'  위 학생은 ( {input_data[4]} )(으)로 인하여 학교성적관리규정 제32조 ( {input_data[3]} )항에 의거 서류를 첨부하여 ( {input_data[5]} )부터 ( {input_data[6]} )까지 ( {input_data[7]} )일간 출석으로 처리하고자 합니다.')
    run_c = paragraph_c.add_run(str_temp)
    run_c.font.size = Pt(12)  

    # 글자 입력하기 / 가운데 정렬 , 글자크기 12
    paragraph_d = doc.add_paragraph()
    paragraph_d.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    str_temp = (f'{input_data[8]}')
    run_d = paragraph_d.add_run(str_temp)
    run_d.font.size = Pt(12)

    # 글자 입력하기 / 오른쪽 정렬 , 글자크기 12
    paragraph_d = doc.add_paragraph()
    paragraph_d.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    str_temp = (f'담 임 :  {input_data[9]}  (인)\n')
    run_d = paragraph_d.add_run(str_temp)
    run_d.font.size = Pt(12)
    
    # 글자 입력하기 / 가운데 정렬 , 글자크기 15
    paragraph_e = doc.add_paragraph()
    paragraph_e.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    str_temp = '군산제일중학교장  귀하'
    run_e = paragraph_e.add_run(str_temp)
    run_e.font.size = Pt(15)    
 
    # 글 입력하기 /  줄 간격 0.6, 글자크기를 10
    paragraph_f = doc.add_paragraph()
    set_paragraph_spacing_line(paragraph_f, 1)
    str_temp =   ('※ 학교성적관리규정 제32조'
             + '\n 교칙에 의거, 출석하여야 할 날짜에 출석하지 않았을 때에는 결석으로 처리한다. 다만, 다음 각항의 하나에 해당되어 학교장이 부득이하다고 인정하거나 또는 허가한 경우에는 출석으로 처리한다.'
             + '\n ① 천재지변, 전염병 등 불가항력의 사유로 인하여 출석하지 못한 경우'
             + '\n ② 병역 관계 등 공적 의무 또는 공권력의 행사로 인하여 출석하지 못한 경우'
             + '\n ③ 학교를 대표한 경기, 경연 대회 참가 및 현장 실습, 훈련 참가, 교환학습, 교외체험학습 등으로 인하여 출석하지 못한 경우'
             + '\n ④ 경조사로 인하여 출석하지 못한 경우'
             + '\n ⑤ 초·중등교육법시행령 제31조(학생의 징계 등) 제1항의 규정에 의한 학교 내의 봉사, 사회봉사, 특별교육이수 기간'
             + '\n ⑥ 생리통이 극심해 출석이 어려운 여학생으로 확인된 경우 (매월 1일에 한하여 출석으로 인정)'
             + '\n ⑦ 건강장애로 선정된 학생의 병원학교 수업 참여 및 교육계획 기간 내에 이수한 화상강의'
             + '\n ⑧ 건강장애 선정 대상자는 아니나 3개월 이상의 치료를 요하는 심각한 부상으로 병원학교수업 참여 또는 화상 강의를 이수하도록 교육청이 허가하고 이에 따라 이수한 화상강의'
             + '\n ⑨ 보호소년 등의 처우에 관한 법률 제 42조의2제1항에 따라 판사 또는 검사가 의뢰한 대안교육 대상 소년이 소년원 또는 소년분류심사원에서 정해진 교육과정을 이수하였을 경우'
             + '\n ⑩ 기타 부득이한 사유로 학교장의 허가를 받아 결석한 경우'
            )
    run_f = paragraph_f.add_run(str_temp)
    run_f.font.size = Pt(10)
    file_name = f'{current_dir}\\특별결석신고서_{input_data[2]}_{input_data[5]}.docx'

    return file_name

##### 생성된 결석신고서를 현재 설정되어있는 기본 프린터로 인쇄하기####
def print_doc_file(file_path) :
    # 워드 파일 열기
    doc = word.Documents.Open(file_path, ReadOnly=True)
    # 워드 파일 인쇄
    doc.PrintOut()
    # 워드 파일 닫기
    doc.Close()
    return True

#####################################################################
####################### 프로그램의 시작 ##############################
#####################################################################

# 결석신고서정보 파일 오픈
if getattr(sys, 'frozen', False):
    # 실행파일로 실행한 경우,해당 파일을 보관한 디렉토리의 full path를 취득
    current_dir = os.path.dirname(os.path.abspath(sys.executable))
else:
    # 파이썬 파일로 실행한 경우,해당 파일을 보관한 디렉토리의 full path를 취득
    current_dir = os.path.dirname(os.path.abspath(__file__))

open_file_name = f'{current_dir}\\결석정보입력.xlsx'
wb = openpyxl.load_workbook(open_file_name, read_only=True)
sh = wb['결석정보']
input_data = ['','','','','','','','','','']


# 출력을 위한 워드 애플리케이션을 열기
word = win32.gencache.EnsureDispatch('Word.Application')

row_i = 2
col_i = 1
# 총 결석신고서 건수를 확인해서 메시지 띄우기 위한 부분
while sh.cell(row_i, 1).value != None:
    row_i = row_i+ 1
messagebox.showinfo("군산제일중학교 결석신고서 자동 출력 프로그램",f'총 출력건수는 {row_i-2}건 입니다. 건당 5초 정도의 시간이 소요됩니다.')


row_i = 2
col_i = 1
# 실제 출력하는 부분
while sh.cell(row_i, 1).value != None:
    
    # 현재 읽어온 행 데이터를 input_data 리스트에 저장
    for i in range(1,10) :
        input_data[i] = str(sh.cell(row_i, i).value)
        input_data[i] = input_data[i].strip()

    print(f'현재 행 데이터 {input_data}')
    

    #  질병결석신고서
    if input_data[3] == '0' :
        # 새로운 결석신고서 docx 문서 생성
        doc = Document()
        file_name_temp = ''
        file_name_temp = Make_Sick_Absence(doc, input_data)
        doc.save(file_name_temp)
        print(f"질병결석신고서 '{file_name_temp}' 파일이 성공적으로 생성되었습니다.")

    #  특별결석신고서
    else :
        # 새로운 결석신고서 docx 문서 생성
        doc = Document()
        file_name_temp = ''
        file_name_temp = Make_Spec_Absence(doc, input_data)
        doc.save(file_name_temp)
        print(f"특별결석신고서 '{file_name_temp}' 파일이 성공적으로 생성되었습니다.")
  
    row_i = row_i+ 1
    print_doc_file(file_name_temp)

# 워드 애플리케이션 종료
word.Quit()
messagebox.showinfo("군산제일중학교 결석신고서 자동 출력 프로그램","결석신고서 출력 요청 완료, 출력물을 확인하세요")


